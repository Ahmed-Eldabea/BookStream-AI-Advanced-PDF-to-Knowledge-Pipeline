import google.genai as genai
import fitz  # PyMuPDF
import time
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from google.generativeai.types import RequestOptions

# --- CONFIGURATION ---
API_KEY = "AIzaSyDN3N2B4xO5BjvEATgz4TGOk81UMrHNSbQ"
FILE_PATH = "Ophthalmology-book.pdf"  # تأكد من أن هذا هو الملف المضغوط
OUTPUT_DOCX = "Ophthalmology_med_summarry.docx"
CHUNK_SIZE = 239  # رفع الحجم لزيادة جودة الربط بين المعلومات
MAX_REQUESTS = 1

# إعداد المكتبة
genai.configure(api_key=API_KEY)
model = genai.GenerativeModel('gemini-2.5-flash')

def setup_document_organization(doc):
    """ضبط الهوامش وتوفير مساحة جانبية للمذاكرة"""
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(2.5) # مساحة للملاحظات اليدوية

def apply_style(run, font_size=12.5, color=RGBColor(47, 54, 64), bold=False):
    """تطبيق التنسيق الجمالي للنصوص"""
    run.font.name = 'Segoe UI'
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    run.bold = bold

def process_markdown_bold(paragraph, text):
    """معالجة النصوص العريضة (Markdown Bold) داخل الفقرات"""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            clean_text = part.replace('**', '')
            run = paragraph.add_run(clean_text)
            apply_style(run, bold=True, color=RGBColor(31, 58, 86))
        else:
            run = paragraph.add_run(part)
            apply_style(run)

def add_formatted_text(doc, text):
    """تحويل رد الذكاء الاصطناعي إلى هيكل وورد منظم ومنسق"""
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line: continue
        
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(10)
        
        # 1. القوائم النقطية
        if line.startswith(('* ', '- ', '+ ')):
            p.style = 'List Bullet'
            process_markdown_bold(p, line[2:].strip())
            
        # 2. العناوين الطبية (Blue Style)
        elif line.startswith('###') or line.startswith('##'):
            p.paragraph_format.space_before = Pt(15)
            run = p.add_run(line.lstrip('#').strip())
            apply_style(run, font_size=15, color=RGBColor(41, 128, 185), bold=True)
        
        # 3. المصطلحات والتعاريف (Purple Style)
        elif ":" in line and len(line.split(":")[0]) < 45:
            parts = line.split(":", 1)
            term_run = p.add_run(parts[0].strip() + ": ")
            apply_style(term_run, bold=True, color=RGBColor(142, 68, 173))
            process_markdown_bold(p, parts[1].strip())
            
        # 4. النصوص العادية
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            process_markdown_bold(p, line)

def start_summarization():
    if not os.path.exists(FILE_PATH):
        print(f"❌ Error: {FILE_PATH} not found!"); return

    pdf_doc = fitz.open(FILE_PATH)
    total_pages = len(pdf_doc)
    pdf_doc.close()
    
    print(f"📘 Deep Study Mode: Professional Synthesis")
    print(f"📄 Total Pages to process: {total_pages}")

    try:
        print(f"📤 Uploading {FILE_PATH}...")
        uploaded_file = genai.upload_file(path=FILE_PATH)
        
        while uploaded_file.state.name == "PROCESSING":
            print(".", end="", flush=True)
            time.sleep(2)
            uploaded_file = genai.get_file(uploaded_file.name)
        print("\n✅ Cloud Context Ready.")
    except Exception as e:
        print(f"\n❌ Upload Error: {e}"); return

    doc = Document()
    setup_document_organization(doc)
    
    # Title Page Section
    main_title = doc.add_heading()
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = main_title.add_run(f'INTEGRATED MEDICAL STUDY GUIDE\n{FILE_PATH.upper()}')
    apply_style(title_run, font_size=26, color=RGBColor(44, 62, 80), bold=True)

    request_count = 0
    for i in range(1, total_pages + 1, CHUNK_SIZE):
        if request_count >= MAX_REQUESTS: break
        
        end_page = min(i + CHUNK_SIZE - 1, total_pages)
        
        # البرومبت الاحترافي المطور
        prompt = f"""
        You are an expert Medical Professor specialize.
        Task: Analyze pages {i} to {end_page} of the provided document.
        Produce a high-yield, cohesive academic summary for medical students.

        Guidelines:
        - INTEGRATION: Summarize the entire range as ONE unit. Avoid page-by-page lists.
        - STRUCTURE: Use professional headings (e.g., Clinical Presentation, Diagnosis, Management).
        - KEY TERMS: Use **bold** for essential medical terminology.
        - FORMAT: Use 'Term: Definition' for jargon and bullet points for classifications.
        - LANGUAGE: Strict Professional English.
        """

        try:
            request_count += 1
            print(f"🚀 Processing Part {request_count} (Pages {i}-{end_page})...")
            
            # التوليد بمهلة زمنية طويلة لضمان الجودة
            response = model.generate_content(
                [uploaded_file, prompt],
                request_options=RequestOptions(timeout=600)
            )
            
            # إضافة ترويسة الجزء في ملف الوورد
            section_h = doc.add_heading(level=1)
            section_run = section_h.add_run(f'Section {request_count}: Essential  Knowledge')
            apply_style(section_run, font_size=18, color=RGBColor(30, 130, 76))
            
            add_formatted_text(doc, response.text)
            doc.save(OUTPUT_DOCX)
            print(f"✅ Part {request_count} completed and saved.")
            
            time.sleep(15) # تأخير بسيط لتجنب الـ Rate Limit
        except Exception as e:
            print(f"❌ Error in Part {request_count}: {e}")
            if "429" in str(e):
                print("🛑 Waiting for Rate Limit reset (60s)...")
                time.sleep(60)

    print(f"\n✨ Mission Completed! Your professional guide is ready.")
    os.startfile(OUTPUT_DOCX)

if __name__ == "__main__":
    start_summarization()