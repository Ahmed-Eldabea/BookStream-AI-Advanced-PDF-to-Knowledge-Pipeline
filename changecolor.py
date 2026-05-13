from docx import Document
from docx.shared import RGBColor

def change_text_color(file_path, output_path, old_rgb, new_rgb):
    """
    تقوم بفتح ملف وورد وتغيير لون أي نص مكتوب باللون القديم إلى اللون الجديد.
    """
    doc = Document(file_path)
    
    # تحويل القيم لسهولة المقارنة
    def colors_match(font_color, target_rgb):
        if font_color is None or font_color.rgb is None:
            return False
        return font_color.rgb == target_rgb

    print(f"🔄 Processing: {file_path}...")
    
    # المرور على كل الفقرات
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # التحقق مما إذا كان لون النص هو الأسود (0, 0, 0)
            if colors_match(run.font.color, old_rgb):
                run.font.color.rgb = new_rgb
                
    # المرور على الجداول أيضاً (لأن كتب ENT غالباً ما تحتوي جداول)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if colors_match(run.font.color, old_rgb):
                            run.font.color.rgb = new_rgb

    doc.save(output_path)
    print(f"✨ Done! Saved to: {output_path}")

# --- التنفيذ ---
old_color = RGBColor(0, 0, 0)      # اللون المراد استبداله
new_color = RGBColor(31, 58, 86)   # اللون الغامق الهادئ الجديد

change_text_color("ENT-summarry.docx", "ENT-summarry2.docx", old_color, new_color)